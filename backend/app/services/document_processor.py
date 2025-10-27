"""
Document processing service for text extraction and chunking.
Handles PDF, DOCX, and image (OCR) documents.
"""

from pathlib import Path
from typing import Optional, List, Dict, Any
import re

# PDF processing
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

# DOCX processing
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# OCR processing
try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None


class DocumentProcessor:
    """
    Service for extracting text from various document formats.
    """

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> tuple[str, int]:
        """
        Extract text from PDF file.

        Returns:
            (extracted_text, page_count)
        """
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is not installed")

        text_parts = []
        page_count = 0

        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)

                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            extracted_text = "\n\n".join(text_parts)
            return extracted_text, page_count

        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file.

        Returns:
            extracted_text
        """
        if DocxDocument is None:
            raise ImportError("python-docx is not installed")

        try:
            doc = DocxDocument(file_path)
            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            extracted_text = "\n\n".join(text_parts)
            return extracted_text

        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")

    @staticmethod
    def extract_text_from_image(file_path: str) -> str:
        """
        Extract text from image using OCR (Tesseract).

        Returns:
            extracted_text
        """
        if Image is None or pytesseract is None:
            raise ImportError("PIL and pytesseract are required for OCR")

        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text

        except Exception as e:
            raise Exception(f"Error extracting text from image: {str(e)}")

    @staticmethod
    def extract_text(file_path: str, file_extension: str) -> tuple[Optional[str], Optional[int]]:
        """
        Extract text from document based on file extension.

        Args:
            file_path: Path to the file
            file_extension: File extension (e.g., '.pdf', '.docx')

        Returns:
            (extracted_text, page_count)
            page_count is None for non-PDF documents
        """
        ext = file_extension.lower()

        try:
            if ext == ".pdf":
                return DocumentProcessor.extract_text_from_pdf(file_path)
            elif ext in [".docx", ".doc"]:
                text = DocumentProcessor.extract_text_from_docx(file_path)
                return text, None
            elif ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
                text = DocumentProcessor.extract_text_from_image(file_path)
                return text, None
            else:
                return None, None

        except Exception as e:
            print(f"Error processing document: {str(e)}")
            return None, None

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean extracted text by removing excessive whitespace and normalizing.
        """
        if not text:
            return ""

        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)

        # Remove multiple newlines (keep max 2)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Strip whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)

        return text.strip()

    @staticmethod
    def chunk_text(
        text: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> List[Dict[str, Any]]:
        """
        Split text into overlapping chunks for RAG indexing.

        Args:
            text: Text to chunk
            chunk_size: Maximum chunk size in characters
            chunk_overlap: Overlap between chunks in characters

        Returns:
            List of chunks with metadata:
            [
                {
                    "text": "chunk text...",
                    "chunk_index": 0,
                    "start_char": 0,
                    "end_char": 1000
                }
            ]
        """
        if not text:
            return []

        chunks = []
        text_length = len(text)
        start = 0
        chunk_index = 0

        while start < text_length:
            end = start + chunk_size

            # If this is not the last chunk, try to break at a sentence boundary
            if end < text_length:
                # Look for sentence endings: . ! ? followed by space or newline
                sentence_ends = [
                    text.rfind('. ', start, end),
                    text.rfind('! ', start, end),
                    text.rfind('? ', start, end),
                    text.rfind('.\n', start, end),
                ]
                sentence_end = max(sentence_ends)

                if sentence_end > start:
                    end = sentence_end + 1

            chunk_text = text[start:end].strip()

            if chunk_text:
                chunks.append({
                    "text": chunk_text,
                    "chunk_index": chunk_index,
                    "start_char": start,
                    "end_char": end,
                })
                chunk_index += 1

            # Move start position with overlap
            start = end - chunk_overlap

            # Prevent infinite loop
            if start >= text_length:
                break

        return chunks

    @staticmethod
    def extract_metadata(text: str) -> Dict[str, Any]:
        """
        Extract metadata from document text (references to standards, regulations).

        Returns:
            {
                "regulations": ["Part B", "Part L"],
                "standards": ["BS 9999", "ISO 19650"],
                "document_type_hints": ["fire safety", "structural"]
            }
        """
        metadata = {
            "regulations": [],
            "standards": [],
            "document_type_hints": [],
        }

        if not text:
            return metadata

        text_lower = text.lower()

        # Extract UK Building Regulations
        regulations = [
            "Part A", "Part B", "Part C", "Part D", "Part E", "Part F",
            "Part G", "Part H", "Part J", "Part K", "Part L", "Part M",
            "Part N", "Part O", "Part P", "Part Q", "Part R", "Part S"
        ]
        for reg in regulations:
            if reg.lower() in text_lower:
                metadata["regulations"].append(reg)

        # Extract British Standards
        bs_matches = re.findall(r'BS\s+\d{4,5}', text, re.IGNORECASE)
        metadata["standards"].extend(set(bs_matches))

        # Extract ISO standards
        iso_matches = re.findall(r'ISO\s+\d{4,5}', text, re.IGNORECASE)
        metadata["standards"].extend(set(iso_matches))

        # Detect document type hints
        type_keywords = {
            "fire_safety": ["fire", "sprinkler", "smoke", "alarm", "evacuation", "fire safety"],
            "structural": ["structural", "foundation", "load", "steel", "concrete", "beam"],
            "mechanical": ["hvac", "ventilation", "heating", "cooling", "ductwork"],
            "electrical": ["electrical", "lighting", "power", "circuit", "wiring"],
            "accessibility": ["accessibility", "disabled", "wheelchair", "ramp", "accessible"],
        }

        for doc_type, keywords in type_keywords.items():
            if any(keyword in text_lower for keyword in keywords):
                metadata["document_type_hints"].append(doc_type)

        return metadata
